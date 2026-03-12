# api/v1/draft_docx.py - POST /api/v1/draft-docx - convert markdown to DOCX
import io
import json
import re
from http.server import BaseHTTPRequestHandler
from docx import Document


def _strip_html(s):
    """Remove raw HTML tags so they don't appear in output."""
    if not s:
        return s
    return re.sub(r"<[^>]+>", "", s)


def _add_paragraph_with_inline(doc, text, style=None):
    """Add a paragraph, converting **bold** to bold runs. Strips HTML."""
    text = _strip_html(text or "")
    p = doc.add_paragraph(style=style)
    pos = 0
    while True:
        m = re.search(r"\*\*([^*]+)\*\*", text[pos:])
        if not m:
            if pos < len(text):
                p.add_run(text[pos:])
            break
        start, end = m.start() + pos, m.end() + pos
        if start > pos:
            p.add_run(text[pos:start])
        run = p.add_run(m.group(1))
        run.bold = True
        pos = end
    return p


def _parse_table(lines, start_i):
    """Parse a markdown table starting at start_i. Returns (list of cell rows, next line index)."""
    rows = []
    i = start_i
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped.startswith("|"):
            break
        # Skip separator line |---|---|
        if re.match(r"^\|[\s\-:]+\|", stripped):
            i += 1
            continue
        cells = [c.strip() for c in stripped.split("|") if c.strip() != ""]
        if not cells:
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def _add_table_with_borders(doc, rows):
    """Add a Word table with the given rows and enable all borders."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.rows[ri].cells[ci]
            cell.text = _strip_html(cell_text or "").replace("**", "").strip()
    return table


def _markdown_to_docx(markdown_text):
    """Convert markdown to docx. Handles # ## ### headers, **bold**, tables, lists, horizontal rules."""
    doc = Document()
    lines = (markdown_text or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("### "):
            _add_paragraph_with_inline(doc, stripped[4:].strip(), style="Heading 3")
        elif stripped.startswith("## "):
            _add_paragraph_with_inline(doc, stripped[3:].strip(), style="Heading 2")
        elif stripped.startswith("# "):
            _add_paragraph_with_inline(doc, stripped[2:].strip(), style="Heading 1")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_paragraph_with_inline(doc, stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith("---"):
            doc.add_paragraph("________________________________________")
        elif stripped.startswith("|") and "|" in stripped[1:]:
            rows, i = _parse_table(lines, i)
            _add_table_with_borders(doc, rows)
            continue
        else:
            _add_paragraph_with_inline(doc, stripped)
        i += 1
    return doc


def _handle_post(body):
    markdown = body.get("markdownContent", "")
    if not markdown:
        return 400, None, "application/json", json.dumps({"error": "markdownContent required"}).encode()
    try:
        doc = _markdown_to_docx(markdown)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return 200, buffer.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", None
    except Exception as e:
        return 500, None, "application/json", json.dumps({"error": str(e)}).encode()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(length).decode("utf-8") if length else "{}"
            body = json.loads(raw) if raw.strip() else {}
        except (ValueError, json.JSONDecodeError):
            body = {}
        status, data, content_type, err_body = _handle_post(body)
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        if data:
            self.send_header("Content-Disposition", 'attachment; filename="qualification-draft.docx"')
            self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        if data:
            self.wfile.write(data)
        elif err_body:
            self.wfile.write(err_body)
